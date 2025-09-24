"""
Parsing Fallbacks and Error Handling

This module provides robust error handling and fallback mechanisms for document parsing.
It includes alternative parsing methods when Granite-Docling fails and quality assessment
for partial parsing recovery.
"""

import logging
import tempfile
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
from uuid import uuid4

import PyPDF2
import docx
from pydantic import BaseModel, Field

from agentic_rag.services.content_extraction import ExtractedContent, LayoutElement, ContentType, DocumentStructure
from agentic_rag.services.docling_client import ParseResponse, DocumentMetadata, ParsedContent
from agentic_rag.services.metadata_extraction import EnrichedMetadata, DocumentProperties

logger = logging.getLogger(__name__)


class FallbackMethod(Enum):
    """Available fallback parsing methods."""
    PYPDF2 = "pypdf2"
    PYTHON_DOCX = "python_docx"
    BASIC_TEXT = "basic_text"
    OCR_FALLBACK = "ocr_fallback"


class ParseQuality(Enum):
    """Parse quality levels."""
    EXCELLENT = "excellent"
    GOOD = "good"
    FAIR = "fair"
    POOR = "poor"
    FAILED = "failed"


class FallbackResult(BaseModel):
    """Result from fallback parsing."""
    
    success: bool = Field(..., description="Whether fallback parsing succeeded")
    method_used: FallbackMethod = Field(..., description="Fallback method that was used")
    quality: ParseQuality = Field(..., description="Quality of the fallback parsing")
    content: Optional[ExtractedContent] = Field(None, description="Extracted content if successful")
    error_message: Optional[str] = Field(None, description="Error message if failed")
    processing_time: float = Field(..., description="Processing time in seconds")
    warnings: List[str] = Field(default_factory=list, description="Warnings during processing")


class ParsingError(BaseModel):
    """Represents a parsing error with context."""
    
    error_type: str = Field(..., description="Type of error")
    error_message: str = Field(..., description="Error message")
    stage: str = Field(..., description="Stage where error occurred")
    recoverable: bool = Field(..., description="Whether error is recoverable")
    suggested_fallback: Optional[FallbackMethod] = Field(None, description="Suggested fallback method")


class ParsingFallbackService:
    """Service for handling parsing failures and providing fallback mechanisms."""
    
    def __init__(self):
        self.fallback_methods = {
            FallbackMethod.PYPDF2: self._fallback_pypdf2,
            FallbackMethod.PYTHON_DOCX: self._fallback_python_docx,
            FallbackMethod.BASIC_TEXT: self._fallback_basic_text,
            FallbackMethod.OCR_FALLBACK: self._fallback_ocr
        }
        
        logger.info("Parsing fallback service initialized")
    
    def handle_parsing_failure(
        self,
        file_content: bytes,
        filename: str,
        original_error: Exception,
        attempted_methods: Optional[List[FallbackMethod]] = None
    ) -> FallbackResult:
        """
        Handle parsing failure by attempting fallback methods.
        
        Args:
            file_content: The document file content
            filename: Original filename
            original_error: The original parsing error
            attempted_methods: Methods already attempted (to avoid retrying)
            
        Returns:
            FallbackResult: Result of fallback parsing attempt
        """
        logger.warning(f"Handling parsing failure for {filename}: {str(original_error)}")
        
        # Analyze the error to suggest appropriate fallback
        error_analysis = self._analyze_parsing_error(original_error, filename)
        
        # Determine fallback methods to try
        fallback_methods = self._determine_fallback_methods(filename, error_analysis, attempted_methods)
        
        # Try each fallback method
        for method in fallback_methods:
            logger.info(f"Attempting fallback method {method.value} for {filename}")
            
            try:
                result = self.fallback_methods[method](file_content, filename)
                if result.success:
                    logger.info(f"Fallback method {method.value} succeeded for {filename}")
                    return result
                else:
                    logger.warning(f"Fallback method {method.value} failed for {filename}: {result.error_message}")
            except Exception as e:
                logger.error(f"Fallback method {method.value} raised exception for {filename}: {str(e)}")
        
        # All fallback methods failed
        logger.error(f"All fallback methods failed for {filename}")
        return FallbackResult(
            success=False,
            method_used=FallbackMethod.BASIC_TEXT,  # Default
            quality=ParseQuality.FAILED,
            error_message="All fallback parsing methods failed",
            processing_time=0.0,
            warnings=["All parsing methods exhausted"]
        )
    
    def assess_parse_quality(self, parse_response: ParseResponse, extracted_content: ExtractedContent) -> ParseQuality:
        """
        Assess the quality of a parse result.
        
        Args:
            parse_response: Original parse response
            extracted_content: Extracted content
            
        Returns:
            ParseQuality: Quality assessment
        """
        if not parse_response.success:
            return ParseQuality.FAILED
        
        # Calculate quality score based on various factors
        score = 0.0
        max_score = 0.0
        
        # Content completeness (40% weight)
        if extracted_content.text_content.strip():
            content_score = min(len(extracted_content.text_content) / 1000, 1.0)  # Normalize to 1000 chars
            score += content_score * 0.4
        max_score += 0.4
        
        # Structure detection (30% weight)
        if extracted_content.elements:
            structure_score = min(len(extracted_content.elements) / 10, 1.0)  # Normalize to 10 elements
            score += structure_score * 0.3
        max_score += 0.3
        
        # Table extraction (20% weight)
        if extracted_content.tables:
            table_score = min(len(extracted_content.tables) / 5, 1.0)  # Normalize to 5 tables
            score += table_score * 0.2
        max_score += 0.2
        
        # Processing success (10% weight)
        if parse_response.pages_processed > 0:
            processing_score = min(parse_response.pages_processed / parse_response.metadata.page_count, 1.0)
            score += processing_score * 0.1
        max_score += 0.1
        
        # Calculate final quality ratio
        quality_ratio = score / max_score if max_score > 0 else 0.0
        
        # Map to quality levels
        if quality_ratio >= 0.9:
            return ParseQuality.EXCELLENT
        elif quality_ratio >= 0.7:
            return ParseQuality.GOOD
        elif quality_ratio >= 0.5:
            return ParseQuality.FAIR
        elif quality_ratio >= 0.2:
            return ParseQuality.POOR
        else:
            return ParseQuality.FAILED
    
    def recover_partial_parsing(
        self,
        partial_content: ExtractedContent,
        file_content: bytes,
        filename: str
    ) -> ExtractedContent:
        """
        Attempt to recover from partial parsing by filling in missing content.
        
        Args:
            partial_content: Partially extracted content
            file_content: Original file content
            filename: Original filename
            
        Returns:
            ExtractedContent: Enhanced content with recovered parts
        """
        logger.info(f"Attempting partial parsing recovery for {filename}")
        
        # Try to extract missing text using fallback methods
        if not partial_content.text_content.strip():
            logger.info("Attempting to recover missing text content")
            fallback_result = self._fallback_basic_text(file_content, filename)
            if fallback_result.success and fallback_result.content:
                partial_content.text_content = fallback_result.content.text_content
                partial_content.elements.extend(fallback_result.content.elements)
        
        # Try to recover missing tables
        if not partial_content.tables:
            logger.info("Attempting to recover missing table content")
            # Could implement table-specific fallback here
        
        # Update processing metadata
        partial_content.processing_metadata["recovery_attempted"] = True
        partial_content.processing_metadata["recovery_timestamp"] = str(uuid4())
        
        return partial_content
    
    def _analyze_parsing_error(self, error: Exception, filename: str) -> ParsingError:
        """Analyze parsing error to determine appropriate fallback strategy."""
        error_message = str(error)
        error_type = type(error).__name__
        
        # Determine if error is recoverable and suggest fallback
        recoverable = True
        suggested_fallback = None
        
        if "timeout" in error_message.lower():
            suggested_fallback = FallbackMethod.BASIC_TEXT
        elif "memory" in error_message.lower():
            suggested_fallback = FallbackMethod.PYPDF2
        elif "format" in error_message.lower() or "corrupt" in error_message.lower():
            suggested_fallback = FallbackMethod.BASIC_TEXT
        elif filename.lower().endswith('.pdf'):
            suggested_fallback = FallbackMethod.PYPDF2
        elif filename.lower().endswith(('.docx', '.doc')):
            suggested_fallback = FallbackMethod.PYTHON_DOCX
        else:
            suggested_fallback = FallbackMethod.BASIC_TEXT
        
        return ParsingError(
            error_type=error_type,
            error_message=error_message,
            stage="primary_parsing",
            recoverable=recoverable,
            suggested_fallback=suggested_fallback
        )
    
    def _determine_fallback_methods(
        self,
        filename: str,
        error_analysis: ParsingError,
        attempted_methods: Optional[List[FallbackMethod]] = None
    ) -> List[FallbackMethod]:
        """Determine which fallback methods to try based on file type and error."""
        if attempted_methods is None:
            attempted_methods = []
        
        methods = []
        
        # Add suggested fallback first
        if error_analysis.suggested_fallback and error_analysis.suggested_fallback not in attempted_methods:
            methods.append(error_analysis.suggested_fallback)
        
        # Add file-type specific methods
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf' and FallbackMethod.PYPDF2 not in attempted_methods:
            methods.append(FallbackMethod.PYPDF2)
        elif file_ext in ['.docx', '.doc'] and FallbackMethod.PYTHON_DOCX not in attempted_methods:
            methods.append(FallbackMethod.PYTHON_DOCX)
        
        # Add basic text extraction as last resort
        if FallbackMethod.BASIC_TEXT not in attempted_methods:
            methods.append(FallbackMethod.BASIC_TEXT)
        
        return methods
    
    def _fallback_pypdf2(self, file_content: bytes, filename: str) -> FallbackResult:
        """Fallback parsing using PyPDF2."""
        import time
        start_time = time.time()
        
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf') as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                # Extract text using PyPDF2
                with open(temp_file.name, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    
                    text_parts = []
                    elements = []
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_parts.append(page_text)
                                
                                # Create basic layout element
                                element = LayoutElement(
                                    id=f"pypdf2_page_{page_num}",
                                    type=ContentType.TEXT,
                                    content=page_text.strip(),
                                    page_number=page_num,
                                    confidence=0.7,  # Lower confidence for fallback
                                    metadata={"extraction_method": "pypdf2"}
                                )
                                elements.append(element)
                        except Exception as e:
                            logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    
                    full_text = "\n\n".join(text_parts)
                    
                    if not full_text.strip():
                        return FallbackResult(
                            success=False,
                            method_used=FallbackMethod.PYPDF2,
                            quality=ParseQuality.FAILED,
                            error_message="No text could be extracted",
                            processing_time=time.time() - start_time
                        )
                    
                    # Create extracted content
                    content = ExtractedContent(
                        document_id=str(uuid4()),
                        structure=DocumentStructure(
                            page_count=len(pdf_reader.pages),
                            document_type="pdf"
                        ),
                        elements=elements,
                        tables=[],
                        text_content=full_text,
                        processing_metadata={
                            "extraction_method": "pypdf2_fallback",
                            "pages_processed": len(pdf_reader.pages)
                        }
                    )
                    
                    return FallbackResult(
                        success=True,
                        method_used=FallbackMethod.PYPDF2,
                        quality=ParseQuality.FAIR,
                        content=content,
                        processing_time=time.time() - start_time,
                        warnings=["Limited formatting and structure detection"]
                    )
        
        except Exception as e:
            return FallbackResult(
                success=False,
                method_used=FallbackMethod.PYPDF2,
                quality=ParseQuality.FAILED,
                error_message=f"PyPDF2 fallback failed: {str(e)}",
                processing_time=time.time() - start_time
            )
    
    def _fallback_python_docx(self, file_content: bytes, filename: str) -> FallbackResult:
        """Fallback parsing using python-docx."""
        import time
        start_time = time.time()
        
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                # Extract text using python-docx
                doc = docx.Document(temp_file.name)
                
                text_parts = []
                elements = []
                
                for i, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                        
                        # Determine content type based on style
                        content_type = ContentType.PARAGRAPH
                        if paragraph.style.name.startswith('Heading'):
                            content_type = ContentType.HEADING
                        
                        element = LayoutElement(
                            id=f"docx_para_{i}",
                            type=content_type,
                            content=paragraph.text.strip(),
                            page_number=1,  # DOCX doesn't have clear page boundaries
                            confidence=0.8,
                            metadata={
                                "extraction_method": "python_docx",
                                "style": paragraph.style.name
                            }
                        )
                        elements.append(element)
                
                full_text = "\n\n".join(text_parts)
                
                if not full_text.strip():
                    return FallbackResult(
                        success=False,
                        method_used=FallbackMethod.PYTHON_DOCX,
                        quality=ParseQuality.FAILED,
                        error_message="No text could be extracted",
                        processing_time=time.time() - start_time
                    )
                
                # Create extracted content
                content = ExtractedContent(
                    document_id=str(uuid4()),
                    structure=DocumentStructure(
                        page_count=1,  # Simplified for DOCX
                        document_type="docx"
                    ),
                    elements=elements,
                    tables=[],  # Could extract tables from doc.tables
                    text_content=full_text,
                    processing_metadata={
                        "extraction_method": "python_docx_fallback",
                        "paragraphs_processed": len(doc.paragraphs)
                    }
                )
                
                return FallbackResult(
                    success=True,
                    method_used=FallbackMethod.PYTHON_DOCX,
                    quality=ParseQuality.GOOD,
                    content=content,
                    processing_time=time.time() - start_time,
                    warnings=["Limited table and image extraction"]
                )
        
        except Exception as e:
            return FallbackResult(
                success=False,
                method_used=FallbackMethod.PYTHON_DOCX,
                quality=ParseQuality.FAILED,
                error_message=f"python-docx fallback failed: {str(e)}",
                processing_time=time.time() - start_time
            )
    
    def _fallback_basic_text(self, file_content: bytes, filename: str) -> FallbackResult:
        """Basic text extraction fallback."""
        import time
        start_time = time.time()
        
        try:
            # Try to decode as text
            text = None
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                return FallbackResult(
                    success=False,
                    method_used=FallbackMethod.BASIC_TEXT,
                    quality=ParseQuality.FAILED,
                    error_message="Could not decode file as text",
                    processing_time=time.time() - start_time
                )
            
            # Clean and structure the text
            text = text.strip()
            if not text:
                return FallbackResult(
                    success=False,
                    method_used=FallbackMethod.BASIC_TEXT,
                    quality=ParseQuality.FAILED,
                    error_message="File contains no readable text",
                    processing_time=time.time() - start_time
                )
            
            # Create basic layout element
            element = LayoutElement(
                id="basic_text_content",
                type=ContentType.TEXT,
                content=text,
                page_number=1,
                confidence=0.5,  # Low confidence for basic text
                metadata={"extraction_method": "basic_text"}
            )
            
            # Create extracted content
            content = ExtractedContent(
                document_id=str(uuid4()),
                structure=DocumentStructure(
                    page_count=1,
                    document_type="text"
                ),
                elements=[element],
                tables=[],
                text_content=text,
                processing_metadata={
                    "extraction_method": "basic_text_fallback",
                    "character_count": len(text)
                }
            )
            
            return FallbackResult(
                success=True,
                method_used=FallbackMethod.BASIC_TEXT,
                quality=ParseQuality.POOR,
                content=content,
                processing_time=time.time() - start_time,
                warnings=["No structure or formatting preserved", "Basic text extraction only"]
            )
        
        except Exception as e:
            return FallbackResult(
                success=False,
                method_used=FallbackMethod.BASIC_TEXT,
                quality=ParseQuality.FAILED,
                error_message=f"Basic text fallback failed: {str(e)}",
                processing_time=time.time() - start_time
            )
    
    def _fallback_ocr(self, file_content: bytes, filename: str) -> FallbackResult:
        """OCR fallback for image-based documents."""
        import time
        start_time = time.time()
        
        # This is a placeholder for OCR fallback
        # In a real implementation, you would use libraries like pytesseract or easyocr
        
        return FallbackResult(
            success=False,
            method_used=FallbackMethod.OCR_FALLBACK,
            quality=ParseQuality.FAILED,
            error_message="OCR fallback not implemented",
            processing_time=time.time() - start_time,
            warnings=["OCR fallback requires additional dependencies"]
        )


# Global fallback service instance
_fallback_service: Optional[ParsingFallbackService] = None


def get_parsing_fallback_service() -> ParsingFallbackService:
    """Get or create the global parsing fallback service instance."""
    global _fallback_service
    
    if _fallback_service is None:
        _fallback_service = ParsingFallbackService()
    
    return _fallback_service
